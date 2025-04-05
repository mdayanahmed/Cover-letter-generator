import streamlit as st
import openai
import os
import fitz  # PyMuPDF
from docx import Document
from fpdf import FPDF
import pytesseract
from PIL import Image
import pandas as pd
import matplotlib.pyplot as plt
from datetime import datetime
from io import BytesIO
from github import Github

# ====== Page Config & LinkedIn link ======
st.set_page_config(page_title="AI Career Assistant", layout="wide", page_icon="🧠")
def inject_css(file_name):
    with open(file_name) as f:
        st.markdown(f"<style>{f.read()}</style>", unsafe_allow_html=True)

inject_css("style.css")



hide_github_style = """
    <style>
    #MainMenu {visibility: hidden;}
    footer {visibility: hidden;}
    .css-1v3fvcr.e8zbici2 {display: none;}
    .stDeployButton {display: none;}
    .css-164nlkn {display: none;}
    header:after {
        content: "🔗 Connect on LinkedIn";
        position: fixed;
        top: 10px;
        right: 20px;
        background-color: #0072b1;
        color: white;
        padding: 6px 12px;
        border-radius: 5px;
        font-size: 14px;
        z-index: 100;
    }
    </style>
"""
st.markdown(hide_github_style, unsafe_allow_html=True)

st.title("🚀 AI Cover Letter & Resume Optimizer")

# ====== OpenAI API Key Note ======
with st.expander("ℹ️ What is an OpenAI API Key?"):
    st.markdown("""
        **OpenAI API Key** ek unique password jaisa hota hai jo aapko OpenAI ki services (GPT-3.5 etc.) use karne deta hai.  
        🔑 [Generate your API key here](https://platform.openai.com/account/api-keys)  
        📋 Copy karke neeche daaliye. Apni key secure rakhiye!
    """)

# ====== API Key Input ======
user_api_key = st.text_input("🔐 Enter your OpenAI API Key", type="password")
if not user_api_key:
    st.warning("Please enter your OpenAI API key to proceed.")
    st.stop()
openai.api_key = user_api_key

# ====== Uploads ======
col1, col2 = st.columns(2)
with col1:
    st.subheader("📄 Upload Resume")
    resume_file = st.file_uploader("PDF, DOCX, TXT", type=["pdf", "docx", "txt"])
    resume_text_input = st.text_area("Or paste resume text here", height=200)
with col2:
    st.subheader("🧾 Upload Job Description")
    jd_file = st.file_uploader("PDF, DOCX, TXT, or Image", type=["pdf", "docx", "txt", "png", "jpg", "jpeg"])
    jd_text_input = st.text_area("Or paste JD text here", height=200)

# ====== Extract Text ======
def extract_text_from_file(uploaded_file):
    if uploaded_file is None: return ""
    _, ext = os.path.splitext(uploaded_file.name)
    if ext == ".pdf":
        with fitz.open(stream=uploaded_file.read(), filetype="pdf") as doc:
            return " ".join([page.get_text() for page in doc])
    elif ext == ".docx":
        return "\n".join([para.text for para in Document(uploaded_file).paragraphs])
    elif ext in [".png", ".jpg", ".jpeg"]:
        image = Image.open(uploaded_file)
        return pytesseract.image_to_string(image)
    elif ext == ".txt":
        return uploaded_file.read().decode("utf-8")
    return ""

resume_text = resume_text_input or extract_text_from_file(resume_file)
job_text = jd_text_input or extract_text_from_file(jd_file)

# ====== Smart Dashboard ======
def show_dashboard(improvements_dict):
    df = pd.DataFrame(list(improvements_dict.items()), columns=["Feature", "Improvement (%)"])
    st.subheader("📊 Smart Resume Improvement Dashboard")

    fig1, ax1 = plt.subplots()
    ax1.pie(df["Improvement (%)"], labels=df["Feature"], autopct="%1.1f%%", startangle=90)
    st.pyplot(fig1)

    fig2, ax2 = plt.subplots()
    ax2.bar(df["Feature"], df["Improvement (%)"], color='skyblue')
    ax2.set_ylabel("Improvement (%)")
    ax2.set_title("Resume Feature-Wise Improvement")
    st.pyplot(fig2)

    excel_path = "resume_improvement.xlsx"
    df.to_excel(excel_path, index=False)
    with open(excel_path, "rb") as f:
        st.download_button("⬇️ Download Excel", f, file_name="resume_improvement.xlsx")

    def save_fig(fig, filename):
        buf = BytesIO()
        fig.savefig(buf, format="png")
        with open(filename, "wb") as f:
            f.write(buf.getbuffer())

    save_fig(fig1, "pie_chart.png")
    save_fig(fig2, "bar_chart.png")

    github_token = st.secrets["github"]["token"]
    repo_name = "mdayanahmed/Cover-letter-generator"
    g = Github(github_token)
    repo = g.get_repo(repo_name)

    def upload_to_github(filepath):
        with open(filepath, "rb") as f:
            content = f.read()
        filename = f"{datetime.now().strftime('%Y-%m-%d_%H-%M-%S')}_{os.path.basename(filepath)}"
        repo.create_file(f"analytics/{filename}", f"Add {filename}", content)

    upload_to_github("pie_chart.png")
    upload_to_github("bar_chart.png")
    upload_to_github(excel_path)
    st.success("✅ Charts and Excel uploaded to GitHub!")

# ====== Buttons ======
if resume_text and job_text:
    colA, colB, colC = st.columns(3)

    with colA:
        if st.button("✉️ Generate Cover Letter"):
            with st.spinner("Generating..."):
                prompt = f"Write a professional cover letter based on:\nResume:\n{resume_text}\nJob Description:\n{job_text}"
                response = openai.ChatCompletion.create(
                    model="gpt-3.5-turbo",
                    messages=[{"role": "user", "content": prompt}],
                    temperature=0.7
                )
                cover_letter = response.choices[0].message.content
                st.text_area("📄 Cover Letter Preview", cover_letter, height=300)
                doc = Document()
                doc.add_heading("Cover Letter", 0)
                doc.add_paragraph(cover_letter)
                doc.save("Cover_Letter.docx")
                with open("Cover_Letter.docx", "rb") as file:
                    st.download_button("⬇️ Download Cover Letter", file, "Cover_Letter.docx")

    with colB:
        if st.button("🛠️ Suggest Resume Improvements"):
            with st.spinner("Analyzing..."):
                prompt = f"""You are a resume expert. Based on the resume and job description, suggest area-wise improvement percentages:
Skill Match: 80%
Action Verbs: 60%
Achievements: 70%
ATS Keywords: 75%
Formatting: 65%
Also give suggestions below.

Resume: {resume_text}
Job: {job_text}"""
                response = openai.ChatCompletion.create(
                    model="gpt-3.5-turbo",
                    messages=[{"role": "user", "content": prompt}],
                    temperature=0.5
                )
                content = response.choices[0].message.content
                st.markdown("### 📈 Resume Suggestions")
                st.markdown(content)

                improvements = {}
                for line in content.splitlines():
                    if ":" in line and "%" in line:
                        key, val = line.split(":")
                        try:
                            improvements[key.strip()] = int(val.strip().replace("%", ""))
                        except:
                            pass
                if improvements:
                    show_dashboard(improvements)

    with colC:
        if st.button("💼 Suggest Job Roles"):
            with st.spinner("Finding roles..."):
                prompt = f"Suggest 5 job roles based on this resume:\n{resume_text}"
                response = openai.ChatCompletion.create(
                    model="gpt-3.5-turbo",
                    messages=[{"role": "user", "content": prompt}]
                )
                st.markdown("### 🧑‍💼 Job Role Suggestions")
                st.markdown(response.choices[0].message.content)

    # ATS & Interview (row 2)
    col4, col5 = st.columns(2)
    with col4:
        if st.button("📊 ATS Match Score"):
            with st.spinner("Checking..."):
                prompt = f"""Act as ATS engine. Compare resume and job description. Return:
- Match % (0-100)
- Matching & Missing Keywords
- Short feedback

Resume: {resume_text}
JD: {job_text}"""
                response = openai.ChatCompletion.create(
                    model="gpt-3.5-turbo",
                    messages=[{"role": "user", "content": prompt}],
                    temperature=0.4
                )
                st.markdown("### 📋 ATS Match Report")
                st.markdown(response.choices[0].message.content)

    with col5:
        if st.button("🎤 Interview Questions"):
            with st.spinner("Generating..."):
                prompt = f"Generate 5 interview questions from this JD:\n{job_text}"
                response = openai.ChatCompletion.create(
                    model="gpt-3.5-turbo",
                    messages=[{"role": "user", "content": prompt}],
                    temperature=0.6
                )
                st.markdown("### ❓ Likely Interview Questions")
                st.markdown(response.choices[0].message.content)

else:
    st.info("📥 Please upload or paste both Resume and Job Description to begin.")
